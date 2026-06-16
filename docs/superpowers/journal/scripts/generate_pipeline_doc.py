"""Génère le document Word décrivant le pipeline Marché → Trade.

Usage : python generate_pipeline_doc.py
Output : C:/Users/xav91/Scalping/scalping/docs/superpowers/journal/2026-06-16-pipeline-marche-vers-trade.docx
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUT = Path(__file__).parent.parent / "2026-06-16-pipeline-marche-vers-trade.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    if size:
        r.font.size = Pt(size)
    return p


def add_code(doc, code):
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, header, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return t


def main():
    doc = Document()
    # Style de base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ──────────────────────────────────────────────────────────────
    # PAGE DE GARDE
    # ──────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Pipeline Marché → Trade")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Scalping Radar — 18 étapes décisionnelles et structurelles")
    r.italic = True
    r.font.size = Pt(14)

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Du tick broker à la position fermée — chaque étape, chaque filtre")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Document généré le 16 juin 2026")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────
    # INTRO
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "Vue d'ensemble", 1)

    add_para(doc,
        "Le système Scalping Radar transforme des mouvements de marché en trades exécutés "
        "à travers 18 étapes successives. À chaque étape, des filtres décisionnels peuvent "
        "écarter un signal. Ce document détaille chacune de ces étapes — son input, sa "
        "logique, son output, et les filtres qui peuvent y tuer un trade.")

    add_para(doc, "Résumé en une phrase :", bold=True)
    add_para(doc,
        "Une bougie 5min arrive (étape 1) → 11 filtres séquentiels (étapes 2-12) → broker valide "
        "(étape 14) → SL/TP auto-clampés (étape 15) → monitor protège (étape 16) → trade clôture "
        "(étape 17) → boucle de feedback (étape 18).",
        italic=True)

    add_para(doc, "")
    add_para(doc,
        "Une bougie sur 200+ génère un push effectif. Une push sur ~30 devient un trade fermé "
        "positif. C'est cette profondeur de filtres et le risk management côté bridge qui font "
        "que des risques théoriques cumulés de €600 deviennent en pratique des pertes réelles de €1.")

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 1
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 1 — Ingestion data marché", 1)
    add_para(doc, "Input : prix temps réel & historique (broker → API)", bold=True)
    add_para(doc, "Source : Twelve Data (REST API + WebSocket)")
    add_bullets(doc, [
        "18 paires définies dans WATCHED_PAIRS de .env",
        "Granularité 5min : analyse principale",
        "Granularité 1h : MTF + ATR (volatilité)",
        "Cache : 200s pour 5min, 900s pour 1h",
    ])
    add_para(doc, "Décision : si API indisponible → pair skip ce cycle (no fake data).")
    add_para(doc, "Output : candles_by_pair (OHLC) + price_current.")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 2
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 2 — Cycle radar (toutes les 3 min)", 1)
    add_para(doc, "Trigger : apscheduler interval 180s déclenche run_analysis_cycle()", bold=True)
    add_para(doc, "Pipeline interne (parallèle pour gain de temps) :")
    add_bullets(doc, [
        "Fetch economic_events (Forex Factory JSON)",
        "Pour chaque pair : fetch 5min + 1h candles",
        "Compute volatility (ATR 14 vs baseline 35)",
        "Compute trend (linear regression slope)",
    ])
    add_para(doc, "Output : economic_events, all_candles, h1_candles, volatility_data, trends.")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 3
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 3 — Détection de pattern", 1)
    add_para(doc, "Code : pattern_detector.detect_patterns(candles, pair)", bold=True)
    add_para(doc, "Patterns détectés (algorithme cherche dans les 30 dernières candles 5min) :")
    add_bullets(doc, [
        "pin_bar_up / pin_bar_down — corps petit + mèche longue",
        "engulfing_bullish / engulfing_bearish — bougie englobe la précédente",
        "range_bounce_up / range_bounce_down — rebond sur support/résistance",
        "momentum_up / momentum_down — 3 bougies consécutives même direction",
        "breakout_up / breakout_down — cassure d'un niveau",
    ])
    add_para(doc, "Décision filtrante : pas de pattern → pas de signal, on passe à la pair suivante.")
    add_para(doc, "Output : liste de PatternHit (pattern_type, direction, strength, candle_idx).")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 4
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 4 — Construction du setup", 1)
    add_para(doc, "Code : calculate_trade_setup(pair, pattern, candles)", bold=True)
    add_para(doc, "Construit :")
    add_bullets(doc, [
        "entry_price = open de la bougie en cours",
        "stop_loss = swing high/low récent ± ATR buffer",
        "take_profit_1 = entry ± 2× sl_dist (R:R 1:2)",
        "take_profit_2 = entry ± 3× sl_dist",
        "risk_reward_1, risk_reward_2 calculés",
    ])
    add_para(doc, "Décision filtrante :")
    add_bullets(doc, [
        "Si SL trop loin (>2% du prix) → setup=None",
        "Si R:R < 1.0 → setup=None",
    ])
    add_para(doc, "Output : objet TradeSetup (ou None).")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 5
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 5 — Enrichissement (contexte)", 1)
    add_para(doc, "Code : analysis_engine.enrich_trade_setup(setup, volatility, trend, events)", bold=True)
    add_para(doc, "Ajoute confidence_score (0-100) calculé par multi-facteurs :")
    add_bullets(doc, [
        "Volatilité (ATR ratio)",
        "Tendance alignment",
        "Macro alignment (DXY, gold, SPX correlation)",
        "Pattern strength",
        "Events proximity (-15min si annonce HIGH dans 30min)",
    ])
    add_para(doc, "Ajoute aussi signal_pattern et notes humaines (ex : « tend H1 alignée bullish »).")
    add_para(doc, "Output : setup enrichi avec score + métadonnées.")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 6
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 6 — Coaching verdict (décision finale)", 1)
    add_para(doc, "Code : coaching.compute_verdict(setup, volatility, trend, events, h1_trend)", bold=True)
    add_para(doc, "Calcule 3 listes :")
    add_bullets(doc, [
        "reasons : facteurs favorables (« Score correct (78/100) », « Tendance 1h aligne le signal »…)",
        "warnings : facteurs défavorables (« Conflit MTF », « Session sous-optimale », « Volatilité extrême »…)",
        "blockers : facteurs bloquants (« Marche fermé », « Macro veto active »…)",
    ])

    add_para(doc, "Règle de décision (avec exceptions crypto-weekend et energy) :", bold=True)
    add_code(doc,
        "# Standard (forex, metal, equity_index)\n"
        "if blockers:                       action = SKIP\n"
        "elif score < 75 OR warnings >= 3:  action = SKIP\n"
        "elif warnings >= 1 AND score < 85: action = WAIT\n"
        "else:                              action = TAKE\n\n"
        "# Crypto weekend (depuis 2026-06-14)\n"
        "if blockers:                       SKIP\n"
        "elif score < 65 OR warnings >= 3:  SKIP\n"
        "else:                              TAKE\n\n"
        "# Energy (depuis 2026-06-15)\n"
        "if blockers:                       SKIP\n"
        "elif score < 60 OR warnings >= 3:  SKIP\n"
        "else:                              TAKE")
    add_para(doc, "Output : verdict avec action ∈ {SKIP, WAIT, TAKE}.")
    add_para(doc, "Décision filtrante : SKIP ou WAIT = pas de push. Seul TAKE continue.")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 7
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 7 — Filtre haute confiance", 1)
    add_para(doc, "Code : filter_high_confidence_setups(setups) ligne analysis_engine.py:638", bold=True)
    add_para(doc, "Filtre : ne garde que confidence_score ≥ MIN_CONFIDENCE_SCORE (=25 par défaut, très bas).")
    add_para(doc,
        "Pourquoi 25 et pas 75 ? Le coaching a déjà filtré par TAKE/SKIP. Ce filtre est juste un "
        "dernier garde-fou pour exclure le bruit total.", italic=True)
    add_para(doc, "Output : liste des setups validés à envoyer.")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 8
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 8 — Dispatch sur 2 channels parallèles", 1)
    add_para(doc, "Code : scheduler.py:270-274", bold=True)
    add_code(doc,
        "if all_trade_setups:\n"
        "    asyncio.create_task(telegram_send_setups(all_trade_setups))   # canal user\n"
        "    asyncio.create_task(mt5_bridge_send_setups(all_trade_setups)) # exécution")
    add_para(doc, "Les deux canaux sont indépendants :")
    add_bullets(doc, [
        "Telegram : alerte sur @ScalpingRadar (canal user)",
        "Bridge : ordre auto-exécuté",
    ])

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 9
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 9 — Pré-filtre stars (bridge)", 1)
    add_para(doc, "Code : mt5_bridge.send_setups() ligne 652", bold=True)
    add_code(doc,
        "allowed_pairs = _STAR_PAIRS_SET | _live_extras | _global_extras\n"
        "setups = [s for s in setups if s.pair in allowed_pairs]")
    add_para(doc, "Composition de allowed_pairs :")
    add_bullets(doc, [
        "_STAR_PAIRS_SET = {XAU, XAG, WTI, ETH}/USD (4 stars Phase 4)",
        "_live_extras = {EUR/USD, GBP/USD, USD/JPY} (forex majors Live)",
        "_global_extras = {BTC, SOL, ADA, XRP, LTC, BCH, DOT}/USD (cryptos promues 13/06)",
    ])
    add_para(doc, "Décision filtrante : si pair pas dans cet ensemble → drop silencieux.")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 10
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 10 — Résolution des destinations (multi-tenant)", 1)
    add_para(doc, "Code : bridge_destinations.resolve_destinations(setup)", bold=True)
    add_para(doc, "Pour chaque setup, retourne la liste des destinations à arroser :")
    add_bullets(doc, [
        "admin_legacy : bridge Demo Xavier (Pepperstone, port 8787)",
        "admin_live : bridge Live Xavier (IC Markets, port 8788)",
        "user:N : EA des Premium users (Cédric = user:2, polling backend)",
    ])
    add_para(doc, "Chaque destination a sa propre config : bridge_url, min_confidence, symbol_map, extra_pairs_allowed.")
    add_para(doc, "Output : pour 1 setup → 1 à N destinations.")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 11
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 11 — _check_rejection (par destination)", 1)
    add_para(doc, "Code : mt5_bridge._check_rejection(setup, destination)", bold=True)
    add_para(doc, "Vérifie 10 conditions dans cet ordre :")

    rows = [
        ["1", "pair_admission_state AUTO_EXEC pour (pair, direction)", "_not_admitted"],
        ["2", "MT5_BRIDGE_BLOCKED_PAIRS blocklist", "pair_blocked"],
        ["3", "kill_switch.manual_enabled", "kill_switch"],
        ["4", "kill_switch.rafale_paused_pairs (pair, dir)", "kill_switch_pair_paused"],
        ["5", "Event blackout (annonce économique)", "event_blackout"],
        ["6", "Données simulées", "simulated_data"],
        ["7", "verdict.action == TAKE", "verdict_blocker"],
        ["8", "SL distance suffisante", "sl_too_close"],
        ["9", "confidence_score ≥ dest.min_confidence", "below_confidence"],
        ["10", "Positions ouvertes < cap par pair", "max_positions_per_pair"],
    ]
    add_table(doc, ["#", "Condition", "Reason si échec"], rows, col_widths=[1.0, 10.0, 5.5])

    add_para(doc, "")
    add_para(doc, "Si tout OK → return None (= push autorisé)")
    add_para(doc, "Sinon → log dans signal_rejections table + return reason.")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 12
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 12 — _push_to_destination", 1)
    add_para(doc, "Code : mt5_bridge._push_to_destination(setup, dest)", bold=True)
    add_para(doc, "Build payload :")
    add_code(doc,
        '{\n'
        '    "pair": "ETH/USD",\n'
        '    "direction": "sell",\n'
        '    "entry": 1665.5,\n'
        '    "sl": 1668.1,\n'
        '    "tp": 1660.5,\n'
        '    "sl_dist": 2.6,\n'
        '    "tp_dist": 5.0,\n'
        '    "risk_money": calc_from_sizing,\n'
        '    "comment": "scalping-radar-2026-06-15",\n'
        '    "confidence": 67.2,\n'
        '    "broker_symbol": dest.symbol_map.get(pair)  # ex: "WTI_N6" pour IC Markets\n'
        '}')
    add_para(doc, "HTTP POST vers dest.bridge_url + \"/order\" avec X-API-Key.")
    add_para(doc, "Side effects :")
    add_bullets(doc, [
        "Insert dans mt5_pushes (track tous les pushes)",
        "Si user:N → insert dans mt5_pending_orders (EA poll cette queue)",
    ])

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 13
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 13 — Bridge VPS reçoit /order", 1)
    add_para(doc, "Code : bridge.py:place_order() ligne 1109 (sur le VPS Stockholm)", bold=True)
    add_para(doc, "Validations bridge :")
    add_bullets(doc, [
        "Champs requis présents (pair, direction, entry, sl, tp)",
        "Cohérence SL/TP vs direction (BUY: sl<entry<tp ; SELL: tp<entry<sl)",
        "MT5_BRIDGE_ALLOWED_ASSET_CLASSES contient la classe",
        "ensure_mt5_connected() OK",
        "resolve_symbol(pair) (= mapping local + broker_symbol du payload)",
        "Si risk_money → calc lots via _compute_lots_from_symbol (= risk_money / (sl_dist × tick_value)), clampé entre volume_min et MAX_LOT_PER_CLASS",
        "_check_safety_gates() : MAX_DAILY_LOSS_PCT, MAX_OPEN_POSITIONS, DEDUP_WINDOW_SEC, TRADING_HOURS_UTC",
    ])
    add_para(doc, "Décision filtrante : si un des gates échoue → HTTP 429 blocked.")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 14
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 14 — OrderSend à MT5", 1)
    add_para(doc, "Code : bridge.py:_send_market_order() ligne 381", bold=True)
    add_para(doc, "Construit MqlTradeRequest :")
    add_bullets(doc, [
        "action = TRADE_ACTION_DEAL (market order)",
        "symbol = mt5_symbol (post-mapping)",
        "volume = lots calculés",
        "type = ORDER_TYPE_BUY / SELL",
        "price = tick.ask / bid",
        "sl = 0, tp = 0 si on a sl_dist > 0 (post-fill)",
        "deviation = DEVIATION_POINTS (slippage tolérance)",
        "magic = MAGIC_NUMBER",
        "type_filling = _pick_filling_mode(symbol) (FOK/IOC/RETURN selon broker)",
    ])

    add_para(doc, "mt5.order_send(request) → broker reçoit", italic=True)

    add_para(doc, "Result possible :", bold=True)
    rows = [
        ["10009 (DONE)", "Trade exécuté, ticket retourné"],
        ["10014 (INVALID_VOLUME)", "Broker refuse le lot (volume_min trop élevé pour ce broker)"],
        ["10016 (INVALID_STOPS)", "SL/TP trop proches du prix"],
        ["10027 (AUTO_TRADING_DISABLED)", "Bouton AutoTrading MT5 OFF"],
        ["0 (généric)", "SymbolSelect a échoué (symbole non trouvé)"],
    ]
    add_table(doc, ["retcode", "Signification"], rows, col_widths=[5.0, 11.0])

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 15
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 15 — Application SL/TP post-fill", 1)
    add_para(doc, "Code : bridge.py:_apply_sltp_from_fill() ligne 468", bold=True)

    add_para(doc, "Pourquoi post-fill :", bold=True)
    add_para(doc,
        "Si on envoie SL/TP avec le market order, le prix de fill peut être différent de l'entry "
        "(slippage). En les posant après, on les recalcule depuis le vrai fill_price, ce qui "
        "préserve le R:R intended par le backend.")

    add_para(doc, "Algorithme :")
    add_bullets(doc, [
        "Récupère fill_price (3 fallbacks : result.price, positions_get(ticket).price_open, tick.bid/ask)",
        "Récupère info.trade_stops_level et info.point (forbidden zone broker)",
        "Calcule new_sl = fill ± sl_dist et new_tp = fill ± tp_dist",
        "Clamp anti-Invalid stops : si SL/TP tombe dans la zone interdite (< stops_level points × 1.2 du prix courant), pousse à la limite extérieure",
        "mt5.order_send({action: TRADE_ACTION_SLTP, position: ticket, sl, tp})",
    ])
    add_para(doc, "Output : SL/TP appliqués à la position côté broker.")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 16
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 16 — Position monitor (run en boucle)", 1)
    add_para(doc, "Code : bridge.py task background interval MONITOR_INTERVAL_SEC=5s", bold=True)
    add_para(doc, "Pour chaque position ouverte :")
    add_bullets(doc, [
        "Break-even : si prix a parcouru BREAKEVEN_TRIGGER_PCT=50% du chemin vers TP1, déplace SL à entry + 1 pip",
        "Partial close : si prix touche TP1, ferme PARTIAL_CLOSE_PCT=50% du volume",
        "Trailing stop : sur le restant, suit le prix à TRAIL_DISTANCE_POINTS=150 (≈ 15 pips forex / $1.50 XAU)",
    ])
    add_para(doc,
        "C'est ÇA qui transforme un SL plein de −€15 en une perte réelle de −€0.04 "
        "(on a vu 99.8 % d'économie sur Live aujourd'hui).", italic=True)

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 17
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 17 — Clôture", 1)
    add_para(doc, "Le trade se ferme automatiquement quand :")
    add_bullets(doc, [
        "SL touché → close_reason = SL",
        "TP1/TP2 touché → close_reason = TP1 ou TP2",
        "Trailing déclenche → close_reason = MANUAL (par le monitor)",
        "Daily loss limit → close_reason = KILL",
    ])
    add_para(doc,
        "Bridge sync EC2 : background task /api/sync query bridge /positions et /deals, "
        "insère dans personal_trades (table EC2) avec PnL final.")

    # ──────────────────────────────────────────────────────────────
    # ÉTAPE 18
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "ÉTAPE 18 — Boucle de feedback", 1)
    add_para(doc, "Tables alimentées en continu :", bold=True)
    add_bullets(doc, [
        "personal_trades : tous les trades exécutés avec PnL final",
        "mt5_pushes : toutes les tentatives bridge",
        "signal_rejections : tous les rejets avec raison",
        "pair_admission_state : state machine pair × direction",
        "pair_pnl_regulator : auto-pause si PnL window < -10%",
        "pair_admission_controller : auto-demote AUTO_EXEC → OBSERVED si max_dd < -10%",
    ])
    add_para(doc, "Surveillance externe :", bold=True)
    add_bullets(doc, [
        "Cron wti-auto-pause.py toutes 5 min : 2 SL consécutifs Live → PAUSE",
        "Notifications Telegram infra (alertes critiques)",
        "Notifications Telegram sales (nouveau client payant Stripe)",
    ])

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────
    # SCHÉMA RÉSUMÉ
    # ──────────────────────────────────────────────────────────────
    add_heading(doc, "Synthèse visuelle du flux", 1)
    add_code(doc,
        "┌─────────────────────────────────────────────────────────────────┐\n"
        "│   1. Twelve Data → OHLC + prix temps réel                        │\n"
        "│   2. Cycle radar 3min → fetch + enrich data                      │\n"
        "│   3. Pattern detector → liste de PatternHit                      │\n"
        "│   4. Calculate setup → entry/sl/tp/RR                            │\n"
        "│   5. Enrich → confidence_score + notes                           │\n"
        "│   6. Coaching verdict → SKIP / WAIT / TAKE                       │\n"
        "│   7. filter_high_confidence → score ≥ 25                         │\n"
        "│   8. Dispatch parallèle :                                        │\n"
        "│        ├─ Telegram canal user                                    │\n"
        "│        └─ MT5 bridge ↓                                           │\n"
        "│   9. Pré-filtre stars                                            │\n"
        "│  10. Résolve destinations (admin_legacy, admin_live, user:N)     │\n"
        "│  11. _check_rejection (10 conditions par dest)                   │\n"
        "│  12. POST /order au bridge VPS                                   │\n"
        "│  13. Bridge VPS valide                                           │\n"
        "│  14. mt5.order_send → broker exécute                             │\n"
        "│  15. _apply_sltp_from_fill (clamp anti-Invalid stops)            │\n"
        "│  16. Position monitor 5s : BE + partial + trailing               │\n"
        "│  17. Clôture (SL / TP / MANUAL / KILL)                           │\n"
        "│  18. Feedback DB + monitoring                                    │\n"
        "└─────────────────────────────────────────────────────────────────┘")

    add_para(doc, "")
    add_para(doc, "Ratio observé (Live IC Markets, 15/06) :", bold=True)
    add_bullets(doc, [
        "Cycles radar : ~480 / jour",
        "Setups générés : ~3 500 / jour",
        "Pushes effectifs : 34 / jour (sur 1 destination Live)",
        "Trades positifs : 2 / jour (5.9 %)",
        "PnL net : −€1.24 sur €300 (−0.4 %)",
    ])

    # ──────────────────────────────────────────────────────────────
    # SAUVEGARDE
    # ──────────────────────────────────────────────────────────────
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"✅ Document généré : {OUT}")
    print(f"   Taille : {OUT.stat().st_size // 1024} KB")


if __name__ == "__main__":
    main()
